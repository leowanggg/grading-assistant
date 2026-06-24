"""
DOCX Parser — Extract exam answers from student .docx files.

Structure:
  Cover page (student info) → exam header → 4 major questions,
  each with 5 sub-questions, each having 程序 (program) and 结果 (result) sections
  answered as screenshot images.
"""

from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass, field
from typing import List, Optional, Tuple
from xml.etree import ElementTree as ET

from docx import Document
from docx.image.image import Image as DocxImage
from docx.opc.constants import RELATIONSHIP_TYPE as RT


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class ImageSection:
    """One scorable section: 程序 (program) or 结果 (result)."""
    section_type: str               # "程序" or "结果"
    image_bytes: Optional[bytes] = None
    text_content: str = ""
    is_missing: bool = False


@dataclass
class SubQuestion:
    """Single sub-question (1-5) inside a major question."""
    index: int                      # 1-based
    question_text: str = ""
    program: Optional[ImageSection] = None
    result: Optional[ImageSection] = None


@dataclass
class MajorQuestion:
    """One major question (一 to 四)."""
    index: int                      # 1-based
    title: str = ""
    sub_questions: List[SubQuestion] = field(default_factory=list)


@dataclass
class StudentInfo:
    """Student identity from cover page."""
    student_id: str = ""
    name: str = ""
    class_name: str = ""


@dataclass
class StudentSubmission:
    """Parsed student answer document."""
    filename: str
    student_info: StudentInfo = field(default_factory=StudentInfo)
    major_questions: List[MajorQuestion] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Namespace helpers
# ---------------------------------------------------------------------------

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS  = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
VML_NS = "urn:schemas-microsoft-com:vml"
O_NS   = "urn:schemas-microsoft-com:office:office"

WP_INLINE   = f"{{{WP_NS}}}inline"
WP_ANCHOR   = f"{{{WP_NS}}}anchor"
A_BLIP      = f"{{{A_NS}}}blip"
R_EMBED     = f"{{{R_NS}}}embed"
R_ID        = f"{{{R_NS}}}id"
W_DRAWING   = f"{{{W_NS}}}drawing"
V_IMAGEDATA = f"{{{VML_NS}}}imagedata"
V_RECT      = f"{{{VML_NS}}}rect"
V_SHAPE     = f"{{{VML_NS}}}shape"
O_OLEOBJECT = f"{{{O_NS}}}OLEObject"


# ---------------------------------------------------------------------------
# WMF metafile → PNG conversion (Windows GDI)
# ---------------------------------------------------------------------------

def _wmf_to_png_via_gdi(wmf_data: bytes) -> Optional[bytes]:
    """Convert WMF/EMF metafile bytes to PNG using Windows GDI.

    Returns PNG bytes on success, None on failure.
    Only works on Windows (uses ctypes + gdi32.dll).
    """
    import ctypes
    from ctypes import wintypes, byref, sizeof, c_void_p, c_char_p

    if len(wmf_data) < 8:
        return None

    gdi32 = ctypes.windll.gdi32
    user32 = ctypes.windll.user32

    # Set up proper argument and return types (critical for 64-bit handles)
    gdi32.SetWinMetaFileBits.restype = wintypes.HANDLE
    gdi32.SetWinMetaFileBits.argtypes = [wintypes.UINT, c_char_p, wintypes.HDC, c_void_p]
    gdi32.PlayEnhMetaFile.restype = wintypes.BOOL
    gdi32.PlayEnhMetaFile.argtypes = [wintypes.HDC, wintypes.HANDLE, ctypes.POINTER(wintypes.RECT)]
    gdi32.DeleteEnhMetaFile.restype = wintypes.BOOL
    gdi32.DeleteEnhMetaFile.argtypes = [wintypes.HANDLE]
    gdi32.GetEnhMetaFileBits.restype = wintypes.UINT
    gdi32.GetEnhMetaFileBits.argtypes = [wintypes.HANDLE, wintypes.UINT, c_void_p]

    hdc_ref = user32.GetDC(0)
    if not hdc_ref:
        return None

    hemf = None
    hbmp = None
    hdc_mem = None
    w_px = h_px = 0

    try:
        # Convert WMF to Enhanced Metafile
        n_size = len(wmf_data)
        buf = ctypes.create_string_buffer(wmf_data, n_size)
        hemf = gdi32.SetWinMetaFileBits(n_size, buf, hdc_ref, None)
        if not hemf:
            return None

        # Get EMF header to determine bounding rectangle
        emf_size = gdi32.GetEnhMetaFileBits(hemf, 0, None)
        if emf_size > 0:
            emf_buf = ctypes.create_string_buffer(emf_size)
            written = gdi32.GetEnhMetaFileBits(hemf, emf_size, emf_buf)
            if written > 0:
                # Parse ENHMETAHEADER: iType(4) nSize(4) rclBounds(16) ...
                import struct
                left, top, right, bottom = struct.unpack_from('iiii', emf_buf.raw, 8)
                w_dev = max(right - left, 1)
                h_dev = max(bottom - top, 1)
                # Convert 0.01mm device units to pixels (~96 DPI)
                w_px = int(w_dev / 26.46 + 0.5)
                h_px = int(h_dev / 26.46 + 0.5)

        # Ensure reasonable minimum dimensions (typical code screenshots)
        if emf_size == 0 or w_px < 400 or h_px < 300:
            w_px, h_px = 1600, 1000

        # Create memory DC & bitmap
        hdc_mem = gdi32.CreateCompatibleDC(hdc_ref)
        hbmp = gdi32.CreateCompatibleBitmap(hdc_ref, w_px, h_px)
        gdi32.SelectObject(hdc_mem, hbmp)

        # White background (FillRect is in user32.dll)
        brush = gdi32.CreateSolidBrush(0x00FFFFFF)
        rect = wintypes.RECT(0, 0, w_px, h_px)
        user32.FillRect(hdc_mem, byref(rect), brush)
        gdi32.DeleteObject(brush)

        # Render EMF into the DC
        emf_rect = wintypes.RECT(0, 0, w_px, h_px)
        if not gdi32.PlayEnhMetaFile(hdc_mem, hemf, byref(emf_rect)):
            return None

        # Extract pixel data via GetDIBits
        class BITMAPINFOHEADER(ctypes.Structure):
            _fields_ = [
                ('biSize', wintypes.DWORD),
                ('biWidth', wintypes.LONG),
                ('biHeight', wintypes.LONG),
                ('biPlanes', wintypes.WORD),
                ('biBitCount', wintypes.WORD),
                ('biCompression', wintypes.DWORD),
                ('biSizeImage', wintypes.DWORD),
                ('biXPelsPerMeter', wintypes.LONG),
                ('biYPelsPerMeter', wintypes.LONG),
                ('biClrUsed', wintypes.DWORD),
                ('biClrImportant', wintypes.DWORD),
            ]

        class BITMAPINFO(ctypes.Structure):
            _fields_ = [('bmiHeader', BITMAPINFOHEADER)]

        bi = BITMAPINFO()
        bi.bmiHeader.biSize = sizeof(BITMAPINFOHEADER)
        bi.bmiHeader.biWidth = w_px
        bi.bmiHeader.biHeight = -h_px  # negative = top-down DIB
        bi.bmiHeader.biPlanes = 1
        bi.bmiHeader.biBitCount = 32
        bi.bmiHeader.biCompression = 0  # BI_RGB

        row_size = ((w_px * 32 + 31) // 32) * 4
        pixels = ctypes.create_string_buffer(row_size * h_px)
        copied = gdi32.GetDIBits(
            hdc_mem, hbmp, 0, h_px, pixels, byref(bi), 0  # 0 = DIB_RGB_COLORS
        )
        if copied == 0:
            return None

        # Build PIL Image from BGRA raw data → PNG
        # GetDIBits returns BGRA with alpha byte undefined (typically 0).
        # Fix alpha to 255 so the image is fully opaque.
        pixels_ba = bytearray(pixels.raw)
        pixels_ba[3::4] = b'\xff' * (len(pixels_ba) // 4)
        from PIL import Image as PILImage
        img = PILImage.frombuffer('RGBA', (w_px, h_px), bytes(pixels_ba), 'raw', 'BGRA', row_size, 1)
        png_buf = io.BytesIO()
        img.save(png_buf, 'PNG')
        return png_buf.getvalue()

    except Exception:
        return None
    finally:
        if hemf:
            gdi32.DeleteEnhMetaFile(hemf)
        if hbmp:
            gdi32.DeleteObject(hbmp)
        if hdc_mem:
            gdi32.DeleteDC(hdc_mem)
        user32.ReleaseDC(0, hdc_ref)


def _is_wmf(data: bytes) -> bool:
    """Check if byte data looks like a WMF/EMF metafile."""
    if len(data) < 4:
        return False
    # Aldus Placeable WMF header
    if data[:4] == b'\xd7\xcd\xc6\x9a':
        return True
    # Standard WMF META_HEADER (type=1 memory, type=2 disk)
    if data[:2] in (b'\x01\x00', b'\x02\x00'):
        return True
    # EMF header
    if data[:4] == b'\x01\x00\x00\x00':
        return True
    return False


def _ensure_renderable(data: bytes) -> bytes:
    """If data is a WMF/EMF metafile, convert to PNG; otherwise return as-is."""
    if _is_wmf(data):
        png = _wmf_to_png_via_gdi(data)
        if png is not None:
            return png
    return data


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def get_image_bytes(para, document: Document) -> Optional[bytes]:
    """Extract the first image from a paragraph, return raw bytes.

    Supports three embedding formats:
    1. Standard inline/floating images via w:drawing > wp:inline > a:blip
    2. VML-based images via v:imagedata (legacy Word format, e.g. WMF metafiles)
    """
    element = para._element

    # ---- Approach 1: Standard DrawingML inline/floating images ----
    container = element.find(f".//{WP_INLINE}")
    if container is None:
        container = element.find(f".//{WP_ANCHOR}")
    if container is None:
        drawing = element.find(f".//{W_DRAWING}")
        if drawing is not None:
            container = drawing.find(f".//{WP_INLINE}")
    if container is not None:
        blip = container.find(f".//{A_BLIP}")
        if blip is not None:
            r_id = blip.get(R_EMBED)
            if r_id is not None:
                try:
                    rel = document.part.rels[r_id]
                    return _ensure_renderable(rel.target_part.blob)
                except (KeyError, AttributeError):
                    pass

    # ---- Approach 2: VML imagedata (legacy WMF/EMF metafile images) ----
    # Look inside v:rect, v:shape, or directly for v:imagedata
    vml_data = element.find(f".//{V_IMAGEDATA}")
    if vml_data is not None:
        # Try r:id (namespaced) first, then bare "id"
        r_id = vml_data.get(R_ID)
        if r_id is not None:
            try:
                rel = document.part.rels[r_id]
                blob = rel.target_part.blob
                # WMF metafiles need conversion to PNG for PIL/Streamlit
                return _ensure_renderable(blob)
            except (KeyError, AttributeError):
                pass

    return None


def has_image(para) -> bool:
    """Quick check if a paragraph contains an image (DrawingML or VML)."""
    element = para._element
    if element.find(f".//{WP_INLINE}") is not None:
        return True
    if element.find(f".//{W_DRAWING}") is not None:
        return True
    if element.find(f".//{V_IMAGEDATA}") is not None:
        return True
    return False


# ---------------------------------------------------------------------------
# Text helpers
# ---------------------------------------------------------------------------

def para_text(para) -> str:
    """Normalized paragraph text (spaces stripped)."""
    return para.text.strip() if para.text else ""


# ---------------------------------------------------------------------------
# Student info extraction (cover page)
# ---------------------------------------------------------------------------

def extract_student_info(doc: Document) -> StudentInfo:
    """Extract 学号/姓名/专业班级 from cover page (first ~15 paragraphs)."""
    info = StudentInfo()

    # Strategy 1 — scan paragraphs for label: value patterns
    pat_id    = re.compile(r"学\s*号\s*[：:]\s*(.+)")
    pat_name  = re.compile(r"姓\s*名\s*[：:]\s*(.+)")
    pat_class = re.compile(r"专业班级\s*[：:]\s*(.+)")

    # Also handle underline-only value after label in same paragraph
    pat_id2    = re.compile(r"学\s*号")
    pat_name2  = re.compile(r"姓\s*名")
    pat_class2 = re.compile(r"专业班级")

    for para in doc.paragraphs[:20]:
        text = para_text(para)
        if not text:
            continue

        m = pat_id.search(text)
        if m and not info.student_id:
            info.student_id = m.group(1).strip()
            continue
        m = pat_name.search(text)
        if m and not info.name:
            info.name = m.group(1).strip()
            continue
        m = pat_class.search(text)
        if m and not info.class_name:
            info.class_name = m.group(1).strip()
            continue

        # Second pass — look for label and grab last non-empty run
        runs = [r.text for r in para.runs if r.text and r.text.strip()]
        joined = "".join(runs).strip()
        if pat_id2.search(text) and not info.student_id:
            # Value is everything after the label
            idx = text.find("号")
            if idx >= 0:
                candidate = text[idx+1:].strip().lstrip("：:").strip()
                if candidate:
                    info.student_id = candidate
        if pat_name2.search(text) and not info.name:
            idx = text.find("名")
            if idx >= 0:
                candidate = text[idx+1:].strip().lstrip("：:").strip()
                if candidate:
                    info.name = candidate
        if pat_class2.search(text) and not info.class_name:
            idx = text.find("班级")
            if idx >= 0:
                candidate = text[idx+2:].strip().lstrip("：:").strip()
                if candidate:
                    info.class_name = candidate

    # Strategy 2 — check first table if no info found
    if not info.student_id or not info.name:
        for table in doc.tables:
            for row in table.rows:
                row_text = "".join(cell.text for cell in row.cells)
                m = pat_id.search(row_text)
                if m:
                    info.student_id = m.group(1).strip()
                m = pat_name.search(row_text)
                if m:
                    info.name = m.group(1).strip()
                m = pat_class.search(row_text)
                if m:
                    info.class_name = m.group(1).strip()
            if info.name:
                break

    # Strategy 3 — regex fallback over all first-page text
    if not info.student_id:
        all_text = " ".join(para_text(p) for p in doc.paragraphs[:20])
        nums = re.findall(r"\b(\d{8,12})\b", all_text)
        if nums:
            info.student_id = nums[0]
    if not info.name:
        all_text = " ".join(para_text(p) for p in doc.paragraphs[:20])
        names = re.findall(r"姓\s*名\s*[：:]?\s*([^\s]{2,4})", all_text)
        if not names:
            names = re.findall(r"([一-鿿]{2,4})", all_text)
            # Filter out common non-name words
            skip = {"实践考核", "须知", "数据挖掘", "课程考核", "专业班级", "任课教师"}
            names = [n for n in names if n not in skip and len(n) >= 2]
        if names:
            info.name = names[0]

    return info


# ---------------------------------------------------------------------------
# Question structure parsing
# ---------------------------------------------------------------------------

_MAJOR_RE = re.compile(r"^[一二三四]\s*[、.].{4,}")   # only Chinese numerals for major Qs
_SUB_TITLE_RE = re.compile(r"^(\d+)\s*[.、]\s*.{2,}")  # "1.生成数据并查看（5分）" — sub-question section titles
_SUB_RE   = re.compile(r"^[（(](\d+)[）)]")           # parenthetical detail markers "(1)", "（1）"
_PROG_RE  = re.compile(r"^程序[：:]")
_RESULT_RE = re.compile(r"^结果[：:]")
_EMPTY_PAT = re.compile(r"^\s*$")


# ---------------------------------------------------------------------------
# Format detection
# ---------------------------------------------------------------------------

def _has_cn_major_headers(doc: Document) -> bool:
    """Pre-scan: does this document use 一、二、三、四 for major questions?"""
    for para in doc.paragraphs:
        text = para_text(para)
        if _MAJOR_RE.search(text):
            return True
    return False


# ---------------------------------------------------------------------------
# Question structure parsing
# ---------------------------------------------------------------------------

_MAJOR_RE = re.compile(r"^[一二三四]\s*[、.].{4,}")   # only Chinese numerals for major Qs
_SUB_TITLE_RE = re.compile(r"^(\d+)\s*[.、]\s*.{2,}")  # "1.生成数据并查看（5分）" — section titles
_SUB_RE   = re.compile(r"^[（(](\d+)[）)]")           # parenthetical detail markers "(1)", "（1）"
_PROG_RE  = re.compile(r"^程序[：:]")
_RESULT_RE = re.compile(r"^结果[：:]")
_EMPTY_PAT = re.compile(r"^\s*$")


def parse_questions(doc: Document, has_cn_major: bool = False) -> List[MajorQuestion]:
    """
    Core state-machine parser.

    Two formats are supported:
      - Format A (has_cn_major=False): "1.XXX" = major Q, "(N)" = sub-Q (e.g. eg1.docx)
      - Format B (has_cn_major=True):  "一.XXX" = major Q, "1.XXX" = sub-Q, "(N)" = detail (e.g. 2_asr.docx)

    States: SEEK_MAJOR → IN_SUB → (on 程序/结果 label) → CAPTURE_IMAGE → back to IN_SUB
    """
    questions: List[MajorQuestion] = []
    cur_q: Optional[MajorQuestion] = None
    cur_sub: Optional[SubQuestion] = None
    state = "SEEK_MAJOR"
    # buffer for sub-question question_text
    sub_text_buffer: List[str] = []
    # track pending section to fill
    pending_section: Optional[str] = None  # "程序" or "结果"

    for para in doc.paragraphs:
        text = para_text(para)

        # ---------- SEEK_MAJOR ----------
        if state == "SEEK_MAJOR":
            if _MAJOR_RE.search(text):
                if len(questions) >= 4:
                    continue
                cur_q = MajorQuestion(index=len(questions) + 1, title=text)
                questions.append(cur_q)
                state = "IN_SUB"
                sub_text_buffer = []
            elif not has_cn_major and _SUB_TITLE_RE.match(text):
                # Format A: "1.XXX" is the major question header
                if len(questions) >= 4:
                    continue
                cur_q = MajorQuestion(index=len(questions) + 1, title=text)
                questions.append(cur_q)
                state = "IN_SUB"
                sub_text_buffer = []
            continue

        # ---------- IN_SUB ----------
        if state == "IN_SUB":
            if _MAJOR_RE.search(text):
                # Format B: Chinese numeral → next major question
                if len(questions) >= 4:
                    continue
                cur_q = MajorQuestion(index=len(questions) + 1, title=text)
                questions.append(cur_q)
                state = "IN_SUB"
                sub_text_buffer = []
                cur_sub = None
                continue

            m_st = _SUB_TITLE_RE.match(text)
            if m_st:
                sub_idx = int(m_st.group(1))
                if has_cn_major:
                    # Format B: "1.XXX" → sub-question under current major
                    cur_sub = SubQuestion(index=sub_idx, question_text=text)
                    cur_q.sub_questions.append(cur_sub)
                else:
                    # Format A: "1.XXX" → next major question
                    if len(questions) >= 4:
                        # Past 4 majors — treat as text
                        if cur_sub is not None:
                            cur_sub.question_text += "\n" + text
                        continue
                    cur_q = MajorQuestion(index=len(questions) + 1, title=text)
                    questions.append(cur_q)
                    sub_text_buffer = []
                    cur_sub = None
                continue

            # Check for sub-question / detail marker "(1)", "（1）"
            m = _SUB_RE.match(text)
            if m:
                sub_idx = int(m.group(1))
                if has_cn_major:
                    # Format B: parenthetical markers are detail text, not sub-questions
                    if cur_sub is not None and text:
                        cur_sub.question_text += "\n" + text
                else:
                    # Format A: "(N)" is the actual sub-question marker
                    cur_sub = SubQuestion(index=sub_idx, question_text=text)
                    cur_q.sub_questions.append(cur_sub)
                continue

            # Check for 程序 label
            if _PROG_RE.match(text):
                img_bytes = get_image_bytes(para, doc)
                if img_bytes is not None:
                    if cur_sub is not None:
                        cur_sub.program = ImageSection(section_type="程序", image_bytes=img_bytes)
                    continue
                pending_section = "程序"
                state = "CAPTURE_IMAGE"
                continue

            # Check for 结果 label
            if _RESULT_RE.match(text):
                img_bytes = get_image_bytes(para, doc)
                if img_bytes is not None:
                    if cur_sub is not None:
                        cur_sub.result = ImageSection(section_type="结果", image_bytes=img_bytes)
                    continue
                pending_section = "结果"
                state = "CAPTURE_IMAGE"
                continue

            # Otherwise accumulate as question description for current sub
            if cur_sub is not None and text:
                cur_sub.question_text += "\n" + text
            continue

        # ---------- CAPTURE_IMAGE ----------
        if state == "CAPTURE_IMAGE":
            # When a section label appears, first check if THIS paragraph
            # already contains an image (common — image in same para as "结果：")
            img_bytes = get_image_bytes(para, doc)

            if _PROG_RE.match(text):
                # Close previous pending section if switching
                if pending_section == "结果" and cur_sub is not None:
                    cur_sub.result = ImageSection(section_type="结果", is_missing=True)
                if img_bytes is not None:
                    if cur_sub is not None:
                        cur_sub.program = ImageSection(section_type="程序", image_bytes=img_bytes)
                    state = "IN_SUB"
                    pending_section = None
                else:
                    pending_section = "程序"
                continue
            if _RESULT_RE.match(text):
                # Close previous pending section if switching
                if pending_section == "程序" and cur_sub is not None:
                    cur_sub.program = ImageSection(section_type="程序", is_missing=True)
                if img_bytes is not None:
                    if cur_sub is not None:
                        cur_sub.result = ImageSection(section_type="结果", image_bytes=img_bytes)
                    state = "IN_SUB"
                    pending_section = None
                else:
                    pending_section = "结果"
                continue

            # Check if a new section/sub-question starts before we captured anything
            new_sub = False
            new_major = False
            m_st = _SUB_TITLE_RE.match(text)
            m_sr = _SUB_RE.match(text) if not m_st else None
            if m_st:
                sub_idx = int(m_st.group(1))
                if has_cn_major:
                    new_sub = True       # Format B: "1.XXX" → new sub-Q
                else:
                    new_major = True     # Format A: "1.XXX" → next major Q
            elif m_sr:
                sub_idx = int(m_sr.group(1))
                if has_cn_major:
                    new_sub = False      # Format B: skip detail marker in capture state
                else:
                    new_sub = True       # Format A: "(N)" → new sub-Q
            # Also check for Chinese-numeral major-question header
            if not new_sub and not new_major and _MAJOR_RE.search(text):
                new_major = True

            if new_sub or new_major:
                # Mark current section as missing, close it
                section = ImageSection(
                    section_type=pending_section, is_missing=True
                )
                if cur_sub is not None:
                    if pending_section == "程序":
                        cur_sub.program = section
                    elif pending_section == "结果":
                        cur_sub.result = section
                pending_section = None

                if new_major:
                    if len(questions) >= 4:
                        continue  # ignore past 4 majors
                    cur_q = MajorQuestion(index=len(questions) + 1, title=text)
                    questions.append(cur_q)
                    cur_sub = None
                else:
                    cur_sub = SubQuestion(index=sub_idx, question_text=text)
                    cur_q.sub_questions.append(cur_sub)
                state = "IN_SUB"
                continue

            # Try to get image from this paragraph
            img_bytes = get_image_bytes(para, doc)
            section = ImageSection(section_type=pending_section)

            if img_bytes is not None:
                section.image_bytes = img_bytes
            elif _EMPTY_PAT.match(text) or not text:
                # Empty paragraph — could be blank line before image
                continue
            elif text:
                section.text_content = text

            if section.image_bytes is None and not section.text_content:
                section.is_missing = True

            # Attach to current sub-question
            if cur_sub is not None:
                if pending_section == "程序":
                    cur_sub.program = section
                elif pending_section == "结果":
                    cur_sub.result = section

            pending_section = None
            state = "IN_SUB"
            continue

    return questions


# ---------------------------------------------------------------------------
# Question signature for variant matching
# ---------------------------------------------------------------------------

def question_signature(questions: List[MajorQuestion]) -> str:
    """Short hash of question titles for variant detection."""
    parts = []
    for q in questions:
        clean = re.sub(r"\s+", "", q.title)[:60]
        parts.append(clean)
    return "|".join(parts)


# ---------------------------------------------------------------------------
# Public parsing API
# ---------------------------------------------------------------------------

def parse_student_docx(file_path: str) -> StudentSubmission:
    """
    Parse a student .docx file.
    Returns StudentSubmission with info, questions, and answer images.
    """
    doc = Document(file_path)
    info = extract_student_info(doc)
    has_cn_major = _has_cn_major_headers(doc)
    questions = parse_questions(doc, has_cn_major=has_cn_major)

    # Ensure 4 major questions with 5 sub-questions each
    for q_idx, q in enumerate(questions):
        # Fill missing sub-questions with placeholders
        existing = {s.index for s in q.sub_questions}
        for i in range(1, 6):
            if i not in existing:
                q.sub_questions.append(
                    SubQuestion(index=i, question_text=f"({i})")

                )
        q.sub_questions.sort(key=lambda s: s.index)

    return StudentSubmission(
        filename=file_path,
        student_info=info,
        major_questions=questions,
    )


def parse_reference_docx(file_path: str) -> List[MajorQuestion]:
    """
    Parse a reference answer .docx.
    Returns list of MajorQuestion (no student info).
    """
    doc = Document(file_path)
    has_cn_major = _has_cn_major_headers(doc)
    return parse_questions(doc, has_cn_major=has_cn_major)


# ---------------------------------------------------------------------------
# Quick test
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys
    path = sys.argv[1] if len(sys.argv) > 1 else r"c:\Users\Silin WANG\Projects\test\ansr.docx"
    sub = parse_student_docx(path)
    print(f"Student: {sub.student_info.name} / {sub.student_info.student_id}")
    print(f"Questions: {len(sub.major_questions)}")
    for q in sub.major_questions:
        print(f"  Q{q.index}: {q.title[:60]}")
        for s in q.sub_questions:
            prog = "✓" if s.program and s.program.image_bytes else ("✗" if s.program and s.program.is_missing else "?")
            res  = "✓" if s.result and s.result.image_bytes else ("✗" if s.result and s.result.is_missing else "?")
            print(f"    ({s.index}) 程序:{prog} 结果:{res}")
